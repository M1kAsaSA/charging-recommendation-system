from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

src = Path(r"c:/charging-recommendation-system/论文/智能充电桩推荐系统_论文初稿.docx")
out = Path(r"c:/charging-recommendation-system/论文/智能充电桩推荐系统_论文初稿_排版.docx")

if not src.exists():
    raise FileNotFoundError(src)

doc = Document(str(src))

# Apply section margins to all sections
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.50)
    section.right_margin = Cm(2.50)
    section.header_distance = Cm(2.30)
    section.footer_distance = Cm(1.80)

# Update paragraph styles based on template info
styles = doc.styles

# Normal body text
normal = styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

# Heading 1 (章节标题)
if "Heading 1" in styles:
    h1 = styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(11.25)

# Heading 2 (节标题)
if "Heading 2" in styles:
    h2 = styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)
    h2.font.bold = False
    h2.paragraph_format.space_before = Pt(11.25)
    h2.paragraph_format.space_after = Pt(11.25)

# Heading 3 (小节标题)
if "Heading 3" in styles:
    h3 = styles["Heading 3"]
    h3.font.name = "宋体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h3.font.size = Pt(10.5)
    h3.font.bold = False
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(13)
    h3.paragraph_format.space_after = Pt(13)

# Set list paragraph font to body style
if "List Paragraph" in styles:
    lp = styles["List Paragraph"]
    lp.font.name = "宋体"
    lp._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    lp.font.size = Pt(12)
    lp.paragraph_format.line_spacing = 1.5

# Find the paragraph that starts the abstract section
abstract_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "摘要":
        abstract_idx = i
        break

if abstract_idx is None:
    raise ValueError("Cannot find '摘要' paragraph for section split")

# Insert a section break before abstract by adding sectPr to previous paragraph
prev_p = doc.paragraphs[abstract_idx - 1] if abstract_idx > 0 else doc.paragraphs[0]
prev_p_pr = prev_p._p.get_or_add_pPr()
sect_pr = OxmlElement("w:sectPr")
prev_p_pr.append(sect_pr)

# Refresh sections after inserting sectPr
sections = doc.sections
if len(sections) < 2:
    # Force add a section if needed
    doc.add_section()
    sections = doc.sections

# Configure section 1 (cover/statement) - no headers/footers
sec1 = sections[0]
sec1.different_first_page_header_footer = True
sec1.header.is_linked_to_previous = False
sec1.footer.is_linked_to_previous = False
sec1.header.paragraphs[0].text = ""
sec1.footer.paragraphs[0].text = ""

# Configure section 2 (main content)
sec2 = sections[1]
sec2.different_first_page_header_footer = False
sec2.odd_and_even_pages_header_footer = True
sec2.header.is_linked_to_previous = False
sec2.footer.is_linked_to_previous = False
sec2.even_page_header.is_linked_to_previous = False
sec2.even_page_footer.is_linked_to_previous = False

# Set header text for odd/even pages
sec2.header.paragraphs[0].text = "河南科技大学毕业设计说明书（论文）"
sec2.even_page_header.paragraphs[0].text = "毕业设计（论文）题目名称"

# Add page number field to footer (centered)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)

# Clear existing footer text and add page numbers
sec2.footer.paragraphs[0].text = ""
add_page_number(sec2.footer.paragraphs[0])
sec2.even_page_footer.paragraphs[0].text = ""
add_page_number(sec2.even_page_footer.paragraphs[0])

# Set page number start at 1 for section 2
sect_pr2 = sec2._sectPr
pg_num_type = sect_pr2.find(qn('w:pgNumType'))
if pg_num_type is None:
    pg_num_type = OxmlElement('w:pgNumType')
    sect_pr2.append(pg_num_type)
pg_num_type.set(qn('w:start'), '1')

# Save output
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f"Wrote {out}")
