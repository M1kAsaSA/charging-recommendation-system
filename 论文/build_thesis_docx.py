from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"c:/charging-recommendation-system/论文")
md_path = root / "智能充电桩推荐系统_论文初稿.md"
out_path = root / "智能充电桩推荐系统_论文初稿.docx"

student_name = "李康辉"
college = "信息工程学院"
major = "物联网工程"
advisor = "范庆辉"

# Use current date if present, otherwise fallback to today
now = datetime.now()
date_str = f"{now.year}年{now.month}月{now.day}日"

doc = Document()

# Cover page
p = doc.add_paragraph("HENAN UNIVERSITY OF SCIENCE & TECHNOLOGY")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("毕 业 设 计（论 文）")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("题目     基于 Spring Boot 的智能充电桩推荐系统的设计与实现")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(f"姓   名         {student_name}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(f"学   院        {college}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(f"专   业        {major}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(f"指导教师       {advisor}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(date_str)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Statements
p = doc.add_paragraph()
p.add_run("学位论文写作声明").bold = True
p = doc.add_paragraph(
    "本人郑重声明：所呈交的学位论文，是本人在导师的指导下，独立进行研究工作所取得的成果。"
    "除文中已经注明引用的内容外，本论文不含任何其他个人或集体已经发表或撰写过的作品或成果。"
    "对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本声明的法律结果由本人承担。"
)
p = doc.add_paragraph("论文作者签名：          日期：    年    月    日")

p = doc.add_paragraph()
p.add_run("学位论文使用授权说明").bold = True
p = doc.add_paragraph(
    "本人完全了解河南科技大学关于收集、保存、使用学位论文的规定，即：按照学校要求提交学位论文的印刷本和电子版本；"
    "学校有权保存学位论文的印刷本和电子版，并提供目录检索与阅览服务；"
    "学校可以采用影印、缩印、数字化或其它复制手段保存论文；"
    "在不以赢利为目的的前提下，学校可以将学位论文编入有关数据库,提供网上服务。（保密论文在解密后遵守此规定）"
)
p = doc.add_paragraph("论文作者签名：          导师签名：")
p = doc.add_paragraph("日期：      年    月    日")

doc.add_page_break()

# Convert markdown body
if not md_path.exists():
    raise FileNotFoundError(md_path)

lines = md_path.read_text(encoding="utf-8").splitlines()

in_math_block = False
for raw in lines:
    line = raw.rstrip()
    if line.startswith("# "):
        text = line[2:].strip()
        p = doc.add_paragraph(text)
        p.style = "Title"
        continue
    if line.startswith("## "):
        text = line[3:].strip()
        p = doc.add_paragraph(text)
        p.style = "Heading 1"
        continue
    if line.startswith("### "):
        text = line[4:].strip()
        p = doc.add_paragraph(text)
        p.style = "Heading 2"
        continue
    if line.startswith("#### "):
        text = line[5:].strip()
        p = doc.add_paragraph(text)
        p.style = "Heading 3"
        continue

    if line.strip() == "$$":
        in_math_block = not in_math_block
        continue

    if in_math_block:
        doc.add_paragraph(line)
        continue

    if line.startswith("- "):
        item = line[2:].strip()
        p = doc.add_paragraph(item)
        p.style = "List Bullet"
        continue

    if not line.strip():
        doc.add_paragraph("")
        continue

    doc.add_paragraph(line)

# Save
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(out_path)
print(f"Wrote {out_path}")
