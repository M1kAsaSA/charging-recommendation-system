import os
from pathlib import Path
from docx import Document
from pypdf import PdfReader

root = Path(r"c:/charging-recommendation-system/论文")

files = [
    root / "示例 毕业设计（论文）撰写规范2025(1)(1).docx",
    root / "李康辉_论文目录.docx",
    root / "开题报告-物联网222李康辉.docx",
]

out_files = []

for docx_path in files:
    if not docx_path.exists():
        print(f"Missing docx: {docx_path}")
        continue
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            row_text = [t for t in row_text if t]
            if row_text:
                parts.append("\t".join(row_text))
    out_path = docx_path.with_suffix(".txt")
    out_path.write_text("\n".join(parts), encoding="utf-8")
    out_files.append(out_path)

pdf_path = root / "毕业论文-物联网211-李志涛-v6.pdf"
if pdf_path.exists():
    reader = PdfReader(str(pdf_path))
    pdf_text = []
    for i, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            text = f"[Page {i} extract error: {e}]"
        text = text.strip()
        if text:
            pdf_text.append(text)
    pdf_out = pdf_path.with_suffix(".txt")
    pdf_out.write_text("\n\n".join(pdf_text), encoding="utf-8")
    out_files.append(pdf_out)
else:
    print(f"Missing pdf: {pdf_path}")

print("Converted files:")
for p in out_files:
    print("-", p)
